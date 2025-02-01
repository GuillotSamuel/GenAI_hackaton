import pandas as pd
from docx import Document
from docx.shared import Inches





# def add_dataframe_to_docx(doc, df, title):
#     """ Ajoute un DataFrame sous forme de tableau dans le document Word. """
#     doc.add_heading(title, level=2)  # Ajouter un titre

#     if df is not None and not df.empty:
#         table = doc.add_table(rows=1, cols=len(df.columns))
#         table.style = "Table Grid"

#         # Ajouter les en-têtes
#         hdr_cells = table.rows[0].cells
#         for i, column_name in enumerate(df.columns):
#             hdr_cells[i].text = column_name

#         # Ajouter les lignes
#         for _, row in df.iterrows():
#             row_cells = table.add_row().cells
#             for i, value in enumerate(row):
#                 row_cells[i].text = str(value)

#         doc.add_paragraph()  # Ajouter un espace après le tableau
#     else:
#         doc.add_paragraph("Aucune donnée disponible pour ce tableau.\n")


def create_report(docx_path, csv_files):
    """ Génère un fichier Word contenant tous les CSV sous forme de tableaux. """
    doc = Document()
    doc.add_heading("Rapport de Données CSV", level=1)  # Titre principal

    for title, file_path in csv_files.items():
        df = open_csv(file_path)  # Charger le CSV
        add_dataframe_to_docx(doc, df, title)  # Ajouter au document Word

    doc.save(docx_path)
    print(f"Rapport généré avec succès : {docx_path}")


if __name__ == "__main__":
    docx_path = "data/processed/data_report.docx"
    


    create_report(docx_path)
