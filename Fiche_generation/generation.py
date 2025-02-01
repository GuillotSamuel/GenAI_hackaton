import pandas as pd
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests


# UTILS

def load_json(file_path):
    with open(file_path, 'r') as f:
        return json.load(f)
    

# CALL LLM

# def call_llm_api(query):
#     api_url = "https://api.your-llm-service.aws.com/endpoint"
#     payload = {
#         "query": query
#     }
#     headers = {
#         "Authorization": "Bearer your_api_key"  # Ajoutez votre clé API si nécessaire
#     }
#     response = requests.post(api_url, json=payload, headers=headers)
#     if response.status_code == 200:
#         return response.json()  # Retourne les données JSON si la requête réussit
#     else:
#         print(f"Erreur : {response.status_code}")
#         return {}


def call_llm_api_mock(query): # TEST
    """
    Simulation de l'appel à un LLM, retourne une réponse factice en JSON.
    
    Arguments :
    - query (str) : La requête envoyée.
    
    Retourne :
    - dict : Une réponse simulée en JSON.
    """
    
    print(f"📩 Envoi de la requête au LLM : {query}")  # Affiche la requête pour debug

    # Simulation de réponse en JSON selon la requête
    response_mock = {
        "dirigeants": [
            {
                "nom": "Jean Dupont",
                "fonction": "Maire",
                "date_naissance": "1965-07-21",
                "formation": ["École Nationale d'Administration", "Master en droit public"],
                "carriere": ["Adjoint au maire", "Conseiller municipal"]
            }
        ]
    }

    print("✅ Réponse simulée reçue !")
    return response_mock 


def create_json(file_path, data):
    with open(file_path, 'w') as f:
        json.dump(data, f, indent=4)
        
        
def get_prompt(path):
    with open(path, 'r') as f:
        return f.read()


def get_all_json_files(data_collectivite, data_metropole):
    
    ### COLLECTIVITE ###
    
    # Logo de la collectivité
    prompt = get_prompt("data/prompts/logo_collectivite.txt")
    create_json(data_collectivite['logo_collectivite'], call_llm_api_mock(prompt))

    # Finances de la collectivité
    prompt = get_prompt("data/prompts/finances_collectivite.txt")
    create_json(data_collectivite['finances_collectivite'], call_llm_api_mock(prompt))
    
    # Présentation de la collectivité
    prompt = get_prompt("data/prompts/presentation_collectivite.txt")
    create_json(data_collectivite['presentation_collectivite'], call_llm_api_mock(prompt))
    
    # Projets verts de la collectivité
    prompt = get_prompt("data/prompts/projets_verts_collectivite.txt")
    create_json(data_collectivite['projets_verts_collectivite'], call_llm_api_mock(prompt))
    
    # Projets sociaux de la collectivité
    prompt = get_prompt("data/prompts/projets_sociaux_collectivite.txt")
    create_json(data_collectivite['projets_sociaux_collectivite'], call_llm_api_mock(prompt))
    
    # Représentant de la collectivité
    prompt = get_prompt("data/prompts/representant_collectivite.txt")
    create_json(data_collectivite['representant_collectivite'], call_llm_api_mock(prompt))
    
    # Budget de la collectivité
    prompt = get_prompt("data/prompts/budget_collectivite.txt")
    create_json(data_collectivite['budget_collectivite'], call_llm_api_mock(prompt))
    
    # Type de collectivité
    prompt = get_prompt("data/prompts/type_collectivite.txt")
    create_json(data_collectivite['type_collectivite'], call_llm_api_mock(prompt))
    
    # A une métropole
    prompt = get_prompt("data/prompts/has_a_metropole.txt")
    create_json(data_collectivite['has_a_metropole'], call_llm_api_mock(prompt))
    
    
    ### METROPOLE ###
    
    # Finances de la métropole
    prompt = get_prompt("data/prompts/finances_metropole.txt")
    create_json(data_metropole['finances_metropole'], call_llm_api_mock(prompt))
    
    # Présentation de la métropole
    prompt = get_prompt("data/prompts/presentation_metropole.txt")
    create_json(data_metropole['presentation_metropole'], call_llm_api_mock(prompt))
    
    # Projets verts de la métropole
    prompt = get_prompt("data/prompts/projets_verts_metropole.txt")
    create_json(data_metropole['projets_verts_metropole'], call_llm_api_mock(prompt))
    
    # Projets sociaux de la métropole
    prompt = get_prompt("data/prompts/projets_sociaux_metropole.txt")
    create_json(data_metropole['projets_sociaux_metropole'], call_llm_api_mock(prompt))
    
    # Représentant de la métropole
    prompt = get_prompt("data/prompts/representant_metropole.txt")
    create_json(data_metropole['representant_metropole'], call_llm_api_mock(prompt))
    
    # Budget de la métropole
    prompt = get_prompt("data/prompts/budget_metropole.txt")
    create_json(data_metropole['budget_metropole'], call_llm_api_mock(prompt))
    
        
# DOCUMENT CREATION

def add_cover_page(doc, data_collectivite, data_metropole):
    # Chargement du fichier has_a_metropole
    has_metropole_data = load_json(data_collectivite['has_a_metropole'])
    has_metropole = has_metropole_data.get('has_a_metropole', False)
    
    # Chargement des informations de la collectivité et de la métropole
    collectivity_name = load_json(data_collectivite['presentation_collectivite']).get('name', 'Collectivité Anonyme')
    if has_metropole:
        metropole_name = load_json(data_metropole['presentation_metropole']).get('name', None)
    
    # Ajouter un titre principal avec un style
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run('Fiche Client')
    title_run.bold = True  # Mettre en gras
    title_run.font.size = Pt(24)  # Taille de police 24
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrer le titre
    
    # Ajouter une image (logo)
    doc.add_paragraph()  # Ajouter un saut de ligne
    doc.add_picture("data/img/logo.png", width=Inches(2.5))  # Modifier la largeur de l'image à 2.5 pouces
    doc.add_paragraph()  # Ajouter un autre saut de ligne pour espacer le logo du texte

    # Ajouter le nom de la collectivité
    collectivity_paragraph = doc.add_paragraph()
    collectivity_run = collectivity_paragraph.add_run(f"Collectivité : {collectivity_name}")
    collectivity_run.font.size = Pt(16)  # Taille de la police 16
    collectivity_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Aligner à gauche
    
    # Si la collectivité a une métropole, ajouter son nom
    if has_metropole and metropole_name:
        metropole_paragraph = doc.add_paragraph()
        metropole_run = metropole_paragraph.add_run(f"Métropole : {metropole_name}")
        metropole_run.font.size = Pt(16)  # Taille de la police 16
        metropole_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Aligner à gauche


def create_report(docx_path, data_collectivite, data_metropole):
    doc = Document()
    add_cover_page(doc, data_collectivite, data_metropole)
    
    doc.save(docx_path)


# MAIN

if __name__ == "__main__":
    docx_path = "data/processed/data_report.docx"
    
    data_collectivite = {
        'logo_collectivite': "data/raw/logo_collectivite.json",
        'finances_collectivite': "data/raw/finances_collectivite.json",
        'presentation_collectivite': "data/raw/presentation_collectivite.json",
        'projets_verts_collectivite': "data/raw/projets_verts_collectivite.json",
        'projets_sociaux_collectivite': "data/raw/projets_sociaux_collectivite.json",
        'representant_collectivite': "data/raw/representant_collectivite.json",
        'budget_collectivite': "data/raw/budget_collectivite.json",
        'type_collectivite': "data/raw/type_collectivite.json",
        'has_a_metropole': "data/raw/has_a_metropole.json",
    }
    
    data_metropole = {
        'finances_metropole': "data/raw/finances_metropole.json",
        'presentation_metropole': "data/raw/presentation_metropole.json",
        'projets_verts_metropole': "data/raw/projets_verts_metropole.json",
        'projets_sociaux_metropole': "data/raw/projets_sociaux_metropole.json",
        'representant_metropole': "data/raw/representant_metropole.json",
        'budget_metropole': "data/raw/budget_metropole.json",
    }

    get_all_json_files(data_collectivite, data_metropole)
    create_report(docx_path, data_collectivite, data_metropole)
